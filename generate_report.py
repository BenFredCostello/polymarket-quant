from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMG_PATH = os.path.join(os.path.dirname(__file__), "backtest_results.png")
OUT_PATH = os.path.join(os.path.dirname(__file__), "Polymarket_Quant_Report.docx")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_section(doc, title):
    h = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    return h


def add_subsection(doc, title):
    h = doc.add_heading(title, level=2)
    h.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
    return h


def body(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── Title block ───────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Polymarket Quantitative Trading System")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("System Report & Backtest Analysis  ·  May 2026")
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
sr.italic = True

doc.add_paragraph()

# ── Links ─────────────────────────────────────────────────────────────────────
link_p = doc.add_paragraph()
link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
link_p.add_run("Resources:  ").bold = True
add_hyperlink(link_p, "GitHub Repository", "https://github.com/BenFredCostello/polymarket-quant")
link_p.add_run("   |   ")
add_hyperlink(
    link_p,
    "Live Paper Trading Tracker (Excel)",
    "https://uoa-my.sharepoint.com/:x:/g/personal/bcos422_uoa_auckland_ac_nz/IQCQYSckFmFNRKtTkBGpKX_jAUliZgQS6Dr1PJlIrYcPQ4o?e=6cDlY2",
)
for run in link_p.runs:
    run.font.size = Pt(10.5)

doc.add_paragraph()

# ── Executive Summary ─────────────────────────────────────────────────────────
add_section(doc, "1.  Executive Summary")
body(
    doc,
    "This system applies quantitative finance pricing models to Polymarket, a decentralised binary "
    "prediction market focused on crypto price outcomes. By pricing each market as a binary option "
    "using a Black-Scholes / Merton Jump Diffusion ensemble and comparing model-implied probabilities "
    "to live market prices, the system identifies edges and sizes positions using fractional Kelly "
    "criterion with a correlation adjustment.",
)
body(
    doc,
    "Over a 20-month backtest (August 2024 – April 2026), starting from a $1,000 notional bankroll, "
    "the system achieved a final equity of $2,258.70, representing a 125.9% ROI and an annualised "
    "Sharpe ratio of 1.69 — well above the 0.5–0.8 range typical of equity markets. Of 179 bets "
    "placed with sufficient market data, 167 were winners (93.3% win rate), with a mean model edge "
    "of 2.9% over market prices.",
)
body(
    doc,
    "The strategy is currently paper trading via the linked Excel tracker while practical constraints "
    "(minimum bet size on Polymarket) are evaluated. Key risks include the thin margin between mean "
    "edge and trading fees, limited backtest history, and the structural risk that growing "
    "sophisticated participation in prediction markets erodes the behavioural premium the strategy "
    "relies on.",
)

# ── Project Overview ──────────────────────────────────────────────────────────
add_section(doc, "2.  Project Overview")
body(
    doc,
    "Polymarket is a CLOB (central limit order book) prediction market where participants trade "
    "binary contracts that resolve to $1 (YES) or $0 (NO) on a stated outcome. Contracts are "
    "on-chain (Polygon), with prices representing the crowd-implied probability of each outcome. "
    "This project focuses exclusively on crypto price markets — e.g. 'Will BTC exceed $80,000 by "
    "end of March 2025?' — which are uniquely suited to quantitative pricing because the underlying "
    "asset price process is observable and has established derivative markets.",
)
body(
    doc,
    "The core thesis is that retail participants (who constitute a significant fraction of "
    "Polymarket volume, particularly in crypto markets) exhibit systematic biases — notably a "
    "bullish skew — that create persistent mispricings relative to a well-calibrated quantitative "
    "model. The system exploits this by pricing binary outcomes from first principles and betting "
    "against market consensus when the modelled edge exceeds a minimum threshold.",
)
body(
    doc,
    "The backtest covers 297 resolved contracts. 201 of these had retrievable historical CLOB price "
    "data and are used for full analysis; the remaining 96 lacked market data at the modelled "
    "pricing date and are excluded from bet simulation (though included in calibration statistics).",
)

# ── Methodology ───────────────────────────────────────────────────────────────
add_section(doc, "3.  Methodology")

add_subsection(doc, "3.1  Pricing Model")
body(
    doc,
    "Each Polymarket contract is treated as a European binary (digital) option: it pays $1 if the "
    "underlying crypto price S(T) exceeds a strike K at expiry T, and $0 otherwise. Two models are "
    "run in parallel and their outputs averaged:",
)
bullet(
    doc,
    "Black-Scholes Binary: closed-form price e^{-rT} · N(d2), where d2 is the standard "
    "log-normal drift-adjusted term. Fast and analytically tractable, but assumes continuous "
    "price paths and underprices tail events.",
)
bullet(
    doc,
    "Merton Jump Diffusion (Monte Carlo, 100,000 paths): extends GBM with a Poisson jump "
    "component — dS = (μ − λκ)S dt + σS dW + JS dN — capturing the gap risk and fat tails "
    "endemic to crypto. Jump parameters (λ = 8 yr⁻¹, μⱼ = −4%, σⱼ = 12%) are calibrated from "
    "three years of historical log-returns by identifying outliers beyond 2.5σ. Antithetic "
    "variates are used for variance reduction. The difference between MC and BS prices is "
    "retained as a fat-tail risk premium signal.",
)
body(
    doc,
    "Edge is defined as: edge = model_prob − market_prob. Positive edge triggers a YES bet; "
    "negative edge (market overprices YES) triggers a NO bet.",
)

add_subsection(doc, "3.2  Volatility: Realised vs. Implied")
body(
    doc,
    "Early versions of the system used only realised volatility (30- and 90-day rolling standard "
    "deviation of log-returns from Yahoo Finance). This was replaced with implied volatility from "
    "the Deribit DVOL index as the primary input wherever available, with realised vol as a "
    "fallback. The justification is straightforward: implied volatility reflects the forward-looking "
    "consensus of professional options traders — the same market participants who set the vol "
    "surface that underpins derivatives pricing. Using their estimate, rather than a lagged "
    "historical measure, makes the model's volatility input consistent with how the market itself "
    "prices uncertainty. Realised vol lags regime changes; implied vol incorporates them as they "
    "happen. A ±3-day search window is used to match the nearest available DVOL observation to "
    "the pricing date.",
)

add_subsection(doc, "3.3  Stochastic Pricing Offset")
body(
    doc,
    "Rather than pricing every contract at a fixed point before expiry, the pricing date is sampled "
    "from a lognormal distribution centred approximately 21 days before expiry. This design choice "
    "involves an explicit tradeoff: fixing a pricing date (e.g. always 30 days out) would require "
    "discarding contracts without data at that exact timestamp, reducing sample size. By instead "
    "synthesising a plausible pricing date stochastically, the backtest retains a larger contract "
    "universe and distributes pricing dates across the realistic operating window, mimicking how "
    "bets would actually be placed. The lognormal distribution is appropriate because pricing "
    "activity has a natural lower bound (near zero days before expiry) and a right tail "
    "corresponding to early-entry positions. The tradeoff is that start dates are not directly "
    "observed — they are modelled — which introduces a source of synthetic noise that could "
    "marginally inflate or deflate measured edge versus a fully observed dataset.",
)

add_subsection(doc, "3.4  Bet Sizing")
body(
    doc,
    "Position sizes are determined by fractional Kelly criterion. The full Kelly fraction for a "
    "binary bet is f = (p·b − (1−p)) / b, where b = (1 − market_prob) / market_prob is the "
    "decimal odds. A scaling factor of 0.30 (30% Kelly) is applied, reflecting three sources of "
    "model uncertainty: realised vol lag, jump parameter estimation error, and regime change risk. "
    "A single bet is capped at 15% of bankroll regardless. When BTC/ETH correlation exceeds 0.75 "
    "(which is most of the time), position sizes are further reduced to avoid implicitly "
    "overleveraging the crypto macro factor — two correlated positions are not two independent bets.",
)

add_subsection(doc, "3.5  Edge Thresholds")
body(
    doc,
    "Bets are only placed when edge falls between a minimum (0% in the backtest; tunable in live "
    "trading) and a maximum of 20%. The upper cap deserves explanation: very large apparent edges "
    "(>20%) almost always reflect data artefacts rather than genuine mispricings — near-expiry "
    "deep-OTM contracts where vol is poorly estimated, or markets that have moved sharply on news "
    "not yet reflected in the CLOB snapshot used for pricing. Placing bets in these conditions "
    "risks being on the wrong side of a large, fast-moving market. As shown in Figure 1, the "
    "contracts beyond the 20% threshold (orange diamonds) have a noticeably higher loss rate, "
    "validating the cap empirically.",
)

# ── Backtest Results ──────────────────────────────────────────────────────────
add_section(doc, "4.  Backtest Results")
body(
    doc,
    "Figure 1 below summarises the full backtest output across 297 resolved contracts "
    "(August 2024 – April 2026).",
)

if os.path.exists(IMG_PATH):
    doc.add_picture(IMG_PATH, width=Inches(5.8))
    cap = doc.add_paragraph(
        "Figure 1: Polymarket backtest dashboard. Top: equity curve (teal) vs $1,000 starting "
        "bankroll (dashed), with individual bets overlaid. Bottom-left: edge vs stake scatter — "
        "green = win, red = loss; diamonds = outside edge thresholds. Bottom-right: model "
        "calibration plot for CLOB-matched contracts (n=201). Table: outcome × bet-side "
        "confusion matrix."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True
else:
    body(doc, "[Figure 1: backtest_results.png not found — insert manually]", italic=True)

doc.add_paragraph()

# Stats table
table = doc.add_table(rows=3, cols=4)
table.style = "Table Grid"
headers = ["Metric", "Value", "Metric", "Value"]
values = [
    ["Final Bankroll", "$2,258.70", "Sharpe Ratio", "1.69"],
    ["ROI", "+125.9%", "Win Rate (bets placed)", "93.3%  (167 / 179)"],
    ["Contracts Analysed", "297  (201 with CLOB data)", "Mean Edge", "2.9%"],
]
hdr_row = table.rows[0]
for i, h in enumerate(headers):
    cell = hdr_row.cells[i]
    cell.text = h
    set_cell_bg(cell, "1A3A5C")
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
for r_idx, row_data in enumerate(values):
    row = table.rows[r_idx + 1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            if c_idx % 2 == 0:
                run.bold = True
        if r_idx % 2 == 0:
            set_cell_bg(cell, "EEF4FB")

doc.add_paragraph()

add_subsection(doc, "4.1  Bet Composition & Bullish Bias")
body(
    doc,
    "The confusion matrix (Figure 1, bottom) reveals a pronounced asymmetry: the model placed "
    "163 winning NO bets against just 4 winning YES bets. This is not a model artefact — it "
    "reflects a real and well-documented behavioural pattern in retail prediction markets. "
    "Retail participants systematically overestimate the probability of positive outcomes, "
    "particularly in crypto, a market with a strong narrative of upside (Wolfers & Zitzewitz, "
    "2004). The model prices both sides and places YES bets where justified (9 total), but the "
    "dominant source of edge has been fading the market's bullish consensus. Both bet directions "
    "are retained because the bias is empirical, not structural — if market composition shifts, "
    "the YES side may become the dominant source of edge.",
)

add_subsection(doc, "4.2  Calibration & Brier Score Caveats")
body(
    doc,
    "The calibration plot (Figure 1, bottom-right) shows model probabilities binned against "
    "realised win rates. The model is reasonably calibrated at high-confidence (>0.8) predictions "
    "but shows deviation in the 0.3–0.6 range — the most competitive segment of the market. "
    "A key limitation of the market Brier score (mean squared error against outcome) is that "
    "Polymarket allows trading right up to the moment of resolution. Late-entry traders with "
    "near-certainty information will push market prices to near 0 or 1, making the market appear "
    "well-calibrated even if early-period prices were poor. This means the market's Brier score "
    "is not a reliable benchmark for comparing the model's early-pricing accuracy.",
)
body(
    doc,
    "That said, calibration accuracy is not strictly required for profitability. Because this "
    "is a binary options structure with defined odds, what matters is whether the model is "
    "differently wrong from the market in a way that generates positive expected value — not "
    "whether it is more accurate in absolute terms. Even a model that is slightly less accurate "
    "overall can be profitable if its errors are uncorrelated with the market's errors and it "
    "captures the correct side of the bet.",
)

# ── Risks ─────────────────────────────────────────────────────────────────────
add_section(doc, "5.  Risks")

add_subsection(doc, "5.1  Fee Erosion")
body(
    doc,
    "Polymarket charges approximately 2% in trading fees on winnings. The system's mean edge "
    "of 2.9% over market price is therefore a razor-thin margin before costs. In live trading, "
    "any slippage, partial fills, or periods of below-average edge could render individual "
    "positions unprofitable after fees. This is the most immediate operational risk and argues "
    "for a higher minimum edge threshold in live deployment than the 0% used in the backtest.",
)

add_subsection(doc, "5.2  Market Saturation & Erosion of the Dumb-Money Premium")
body(
    doc,
    "The strategy's fundamental premise is that a meaningful fraction of Polymarket volume comes "
    "from behaviourally biased retail participants. This is plausible today, but prediction "
    "markets are attracting increasing institutional and sophisticated retail attention — "
    "particularly following high-profile political prediction market activity in 2024. As "
    "smart money enters and arbs away persistent biases, the expected edge on any systematic "
    "strategy will compress. Wolfers & Zitzewitz (2004) document how thin mispricings in "
    "mature prediction markets already are; Polymarket's crypto markets are younger and less "
    "efficient, but that gap will likely close. The Excel tracker is monitoring live edge "
    "realisation specifically to detect this compression.",
)

add_subsection(doc, "5.3  Limited Backtest History")
body(
    doc,
    "The backtest spans approximately 20 months and 297 contracts. This is a short history for "
    "drawing statistically robust conclusions — particularly given that the strategy's win rate "
    "is high and the sample is concentrated in a specific market regime (post-2024 crypto bull "
    "run followed by consolidation). Overfitting risk is partially mitigated by avoiding "
    "parameter tuning on the outcome variable and by using physically motivated model "
    "parameters rather than fitted ones, but it cannot be eliminated. Adding simulated market "
    "trends (e.g. bear market periods) to stress-test the backtest is a planned improvement.",
)

add_subsection(doc, "5.4  Correlation & Concentration Risk")
body(
    doc,
    "All current markets are on BTC and ETH, which are highly correlated (ρ > 0.75 most of "
    "the time). A single macro shock — a major exchange failure, regulatory action, or "
    "liquidity event — could move both assets simultaneously and resolve multiple open "
    "positions adversely at once. The fractional Kelly sizing and correlation penalty in the "
    "sizing model are the primary mitigations: the system is designed to never treat two "
    "correlated positions as fully independent bets. Expanding into less correlated assets "
    "would reduce this concentration.",
)

add_subsection(doc, "5.5  Paper Trading Status")
body(
    doc,
    "The live Excel tracker linked above is currently paper trading. Polymarket enforces "
    "minimum bet sizes that make fractional Kelly stakes impractical at the current bankroll "
    "scale. Live deployment is contingent on either scaling the bankroll or identifying "
    "markets where the minimum bet is compatible with the recommended stake. Paper trading "
    "results should be interpreted with this in mind — execution quality, partial fills, and "
    "fee drag are not yet observed in live conditions.",
)

# ── Future Improvements ───────────────────────────────────────────────────────
add_section(doc, "6.  Future Improvements")

add_subsection(doc, "6.1  Expanded Data & Market Coverage")
body(
    doc,
    "The current dataset is limited to ~500 resolved Polymarket crypto contracts. A natural "
    "extension is to source additional historical contracts from alternative data providers "
    "or from Polymarket's on-chain event logs, which could significantly expand backtest depth. "
    "Expanding beyond BTC/ETH to include smaller-cap crypto assets is also worth exploring: "
    "smaller markets tend to have thinner liquidity and more behavioural participation, "
    "potentially offering larger and more persistent edges — though at the cost of wider "
    "bid/ask spreads.",
)

add_subsection(doc, "6.2  Additional Signals from Options Markets")
body(
    doc,
    "The system already extracts implied volatility from Deribit DVOL. The same API provides "
    "richer data: volatility skew, term structure, and put/call ratios. These could be "
    "incorporated as additional features — for example, a steep downside skew might "
    "systematically adjust the jump parameters, or a flat term structure might signal a "
    "low-information regime where the model's edge is less reliable.",
)

add_subsection(doc, "6.3  Cross-Market Correlation & Regime Detection")
body(
    doc,
    "Polymarket crypto prices do not exist in isolation. BTC dominance, DeFi TVL, broader "
    "macro risk indicators (VIX, credit spreads), and on-chain metrics are all correlated "
    "with crypto price outcomes. Cross-referencing Polymarket positions against these signals "
    "could improve timing — for example, avoiding YES bets on BTC upside during periods of "
    "elevated macro risk.",
)

add_subsection(doc, "6.4  Pattern Analysis: Returns by Contract Characteristics")
body(
    doc,
    "A structured statistical analysis (potentially in R) of returns segmented by contract "
    "characteristics would be valuable: bet length (time to expiry at entry), market volume, "
    "asset (BTC vs ETH), and edge magnitude. Preliminary observation suggests the NO-bet "
    "dominance is robust, but it is unclear whether short-dated or long-dated contracts "
    "generate better risk-adjusted returns, and whether there is an optimal volume window "
    "that filters out both illiquid (wide spread) and heavily-traded (efficient) markets. "
    "Modelling these patterns quantitatively would allow tighter bet selection criteria.",
)

add_subsection(doc, "6.5  Comparison to Crypto Market Returns")
body(
    doc,
    "A natural benchmark comparison is the strategy's risk-adjusted return versus simply "
    "holding BTC or ETH over the same period. This comparison is currently limited but "
    "would be instructive: it tests whether the strategy is genuinely generating alpha or "
    "is implicitly long crypto via its bet composition. Given the NO-bet dominance, the "
    "strategy may in fact be net short crypto sentiment, which would make it a useful "
    "diversifier in a crypto-heavy portfolio. Formalising this analysis — including "
    "correlation of strategy returns to BTC daily returns — is a planned next step.",
)

# ── References ─────────────────────────────────────────────────────────────────
add_section(doc, "References")
refs = [
    "Black, F. & Scholes, M. (1973). The Pricing of Options and Corporate Liabilities. "
    "Journal of Political Economy, 81(3), 637–654.",
    "Kelly, J. L. (1956). A New Interpretation of Information Rate. Bell System Technical "
    "Journal, 35(4), 917–926.",
    "Merton, R. C. (1976). Option Pricing When Underlying Stock Returns Are Discontinuous. "
    "Journal of Financial Economics, 3(1–2), 125–144.",
    "Thorp, E. O. (2006). The Kelly Criterion in Blackjack, Sports Betting, and the Stock "
    "Market. In Zenios & Ziemba (Eds.), Handbook of Asset and Liability Management. Elsevier.",
    "Wolfers, J. & Zitzewitz, E. (2004). Prediction Markets. Journal of Economic "
    "Perspectives, 18(2), 107–126.",
]
for r in refs:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(r)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
